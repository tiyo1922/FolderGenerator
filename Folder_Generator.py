import tkinter as tk
from tkinter import filedialog, colorchooser, messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Cm
import os
import datetime

def generate_folders_and_docs():
    try:
        with open(entry_code_path.get(), 'r', encoding='utf-8') as file:
            folder_names = [line.strip() for line in file.readlines() if line.strip()]
    except Exception as e:
        messagebox.showerror("Error", f"Gagal membaca file code.txt: {e}")
        return

    output_dir = folder_output_path.get()
    if not output_dir:
        messagebox.showerror("Error", "Folder output belum dipilih.")
        return

    footer_text = footer_text_area.get("1.0", tk.END).strip()
    footer_color = footer_color_label.cget("fg")
    save_footer_cache(footer_text)

    for folder_name in folder_names:
        folder_path = os.path.join(output_dir, folder_name)
        os.makedirs(folder_path, exist_ok=True)
        doc_name = f"{folder_name}.docx"
        doc_path = os.path.join(folder_path, doc_name)

        doc = Document()

        # Set margin ke 2.54 cm (1 inch)
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Judul Tengah dan Bold
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("UNDANGAN WEB FOTO \nDESAIN NO.\n")
        run.bold = True
        paragraph.alignment = 1  # Center

        # Keterangan Kiri dan Bold
        paragraph2 = doc.add_paragraph()
        run2 = paragraph2.add_run("Urut Nama didahulukan :")
        run2.bold = True
        paragraph2.alignment = 0  # Left
        run2.font.name = 'Cambria'
        run2.font.size = Pt(14)

        # Tambahkan footer resmi (Word footer section)
        if footer_text:
            footer = doc.sections[0].footer
            footer_paragraph = footer.paragraphs[0]
            run_footer = footer_paragraph.add_run(footer_text)
            run_footer.font.name = 'Cambria'
            run_footer.font.size = Pt(12)
            footer_paragraph.alignment = 1  # Center

            # Footer color
            try:
                r, g, b = root.winfo_rgb(footer_color)
                run_footer.font.color.rgb = RGBColor(r >> 8, g >> 8, b >> 8)
            except:
                pass

        doc.save(doc_path)

    log_operation(f"Berhasil membuat {len(folder_names)} folder dan file .docx di {output_dir}")
    messagebox.showinfo("Sukses", "Semua folder dan dokumen berhasil dibuat!")

def browse_code_file():
    path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
    if path:
        entry_code_path.set(path)

def browse_output_folder():
    path = filedialog.askdirectory()
    if path:
        folder_output_path.set(path)

def choose_footer_color():
    color = colorchooser.askcolor()[1]
    if color:
        footer_color_label.config(fg=color)

def save_footer_cache(text):
    with open("footer_cache.txt", "w", encoding='utf-8') as f:
        f.write(text)

def load_footer_cache():
    if os.path.exists("footer_cache.txt"):
        with open("footer_cache.txt", "r", encoding='utf-8') as f:
            return f.read()
    return ""

def log_operation(content):
    with open("log.txt", "a", encoding='utf-8') as f:
        f.write(f"[{datetime.datetime.now()}] {content}\n")

# GUI Setup
root = tk.Tk()
root.title("Folder Generator by AL")
root.geometry("500x500")

# Set icon
try:
    icon_path = os.path.join(os.path.dirname(__file__), "logo.ico")
    root.iconbitmap(icon_path)
except Exception as e:
    print(f"Icon gagal di-load: {e}")

entry_code_path = tk.StringVar()
folder_output_path = tk.StringVar()

tk.Label(root, text="Pilih file code.txt:").pack(pady=5)
tk.Entry(root, textvariable=entry_code_path, width=50).pack()
tk.Button(root, text="Browse", command=browse_code_file).pack(pady=2)

tk.Label(root, text="Pilih folder output:").pack(pady=5)
tk.Entry(root, textvariable=folder_output_path, width=50).pack()
tk.Button(root, text="Browse", command=browse_output_folder).pack(pady=2)

tk.Label(root, text="Teks Footer:").pack(pady=(10, 3))
footer_text_area = tk.Text(root, height=4, width=45)
footer_text_area.insert(tk.END, load_footer_cache())
footer_text_area.pack(pady=3)

tk.Button(root, text="Pilih Warna Footer", command=choose_footer_color).pack(pady=5)
footer_color_label = tk.Label(root, text="Contoh Warna Footer", fg="#000000")
footer_color_label.pack()

tk.Button(root, text="Generate", command=generate_folders_and_docs, bg="navy", fg="white").pack(pady=15)

tk.Label(root, text="PUNAKAWAN INSPIRA INDONESIA", fg="gray").pack(side="bottom", pady=10)

root.mainloop()
